from docx import Document
from docx.shared import Pt
from pathlib import Path

out = Path(r"c:\Users\bilal\Documents\Prj\Portfolio\prj\portfolio.docx")
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def p(text=""):
    doc.add_paragraph(text)

h1("Portfolio")
p("Bilal Drabo")

h2("Navigation")
p("Home")
p("Projects")
p("Soft Skills")
p("Internship")
p("Contact")

h1("Home")
p("Industrial IT student at HELHa Charleroi.")
p("Presentation🙋")
p("I am Bilal, I am 22 years old, a student in industrial informatics at HELHa Charleroi.\nI am actively looking for a traineeship in IT, industrial automation, or web/application development.\nWhen I finished high school, I did not know what I would learn for my future.\nI was very undecided until a friend told me about programming, and I discovered IT.\nI did a lot of research during my holidays and quickly loved it. My first bachelor's was complicated because of a lack of organization. Indeed, I am the eldest son of a family where nobody had gone to a city school, so I didn’t expect the level of organization needed for such graduation.\n")
p("Objective🎯")
p("I would like to find a traineeship in IT, web or mobile programming, in industrial automation on machines such as Schneider or Siemens.")
p("Hobbies⚽")
p("I like sports. I think it’s important for balance with the informatics profession, where we move less. I practise individual sports like fitness or walking, but I also like team sports such as football. In addition I a like long formats such as podcasts or reading.")
p("Video Editing🎥")
p("I have notions of Adobe Suite, such as Premiere Pro and Photoshop, learned by myself.")
p("Automation 🤖")
p("During my studies, I learned to use Schneider and Siemens software. I also learned to make electronic plans and use automation systems through a firewall like Flexy or in a network with eCatcher.")
p("Mobile Development 📱")
p("I can code web applications in React Native or Java.")
p("student job 🧑‍🔧")
p("I was a production worker in my holidays .IT allowed me to see automata working in a professional work environment which helped me greatly for my automate test . I also improved my ability to work in a coordinated team.")

h1("Soft Skills")
p("💻 High-Tech Passion")
p("I am very interested in high-tech and regularly conduct research in this field. \n        I am constantly on the lookout for new technologies, innovations, and industry trends.")
p("📊 Management & Organization")
p("My first bachelor's degree taught me the crucial importance of organization. \n        As the first in my family to pursue higher education, I had to develop my own \n        management and organizational methods to succeed.")
p("🚀 Creativity & Innovation")
p("I enjoy inventing, creating, and finding innovative solutions to technical problems. \n        My journey has taught me to approach challenges with a creative mindset and a willingness to innovate.")
p("💪 Perseverance")
p("Despite initial challenges due to my lack of family academic experience, \n        I persevered and developed strong adaptability. Each obstacle is an opportunity for learning.")
p("🌟 Vision & Ambition")
p("I aim to contribute to the technology field through my work and ideas. \n        My goal is to combine my technical knowledge with creativity to develop \n        solutions that have a positive impact.")

h1("Projects")
projects = [
("Dyce game", "group", "More", "Disciplinary Project", "My initial academic project focused on a dice game that was randomly selected. My team and I had to code an app that randomly throws two dice and displays their values on the computer and an electronic model. As part of our first project, our teachers introduced us to professional organization systems like mind maps, Gantt, and PERT charts. Furthermore, each team had meetings every week to discuss their research, misunderstandings, or improvements. In the middle of the year one of my team left the school so I had to switch beetwen my initial task and other task release by my previous teammates. This challenge has really improved my versatile ability. This project has also helped me in managing shyness because in different meeting we had to speak a lot."),
("Defense Turret", "solo", "More", "Defense Turret", "The third disciplinary project was a solo endeavor. The only requirement was to create an autonomous system where the user could trigger different values with a button. I decided to make a fake defense system. I used a radar sensor mounted on a 360° servo to detect any object in its path, and a solenoid fixed on another 360° servo to shoot an arrow at the target. The system was connected to a React website."),
("Automate Simulation", "solo", "More", "Automate Simulation", "The aim of this automation project was to create a simulation in the Unity Pro XL software that sorts crates on an assembly line. We had three different coloured boxes: grey, green, and blue. When the system is started, a sensor detects the colour of each container and, if it matches, pushes it onto another assembly line, incrementing a counter for the matching crate. If no match is detected, the container moves on to the next sensor or to the end of the line. The line has an external emergency button for safety. We created a GRAFCET diagram and a ladder program, and finally developed an evaluation screen linked to the project."),
("Physic Project Automate", "group", "More", "Physical Project", "The second automation project consisted of wiring an automation system to power. Connecting a button and a lamp. It may seem simple. But it was challenging for beginners. We also had to program a Unity automation system to use two components. My project involved a storage area. When a button was pressed, an assembly line pushed a box into a crane which moved to a selected area. When the area was empty, the crane deposited the box and the lamp light up."),
("Automation with PID and PWM", "group", "More", "Automation with PID and PWM", "As part of this project, we had to turn a heated LED on or off using two buttons (ON and OFF). Then, we regulated the temperature of the case to prevent it from overheating by operating two fans controlled by PWM (Pulse Width Modulation). We added a PID block (Proportional, Integral, Derivative) to control the fans’ speed based on the temperature. Finally, we configured an HMI for real-time visualization and interaction with the system."),
("Industrial Project", "group", "More", "Technocampus All In One Project", "We have to carry out a project that consists of sending data from SIF 401 (PLC) in a secure way using an Ewon Flexy 201 and including different communication protocols. This project will be divided into several stages, the first stage is the analysis and understanding of the hardware used, it will take in hand the Ewon and its configuration with talk2me, to then connect with the ECatcher application, It will also be necessary to understand and use the OPC-UA protocol, which is a communication standard used in the field of industrial automation. The second step is the data communication to the web application, because as mentioned above, we will use the OPC-UA protocol, in order to transmit some data from the SIF401 to our Node-RED application. It will be configured to transmit these OPC-UA data in Web Socket to our web application made in react. The web application will have a login page, which will give the choice between an OAuth authentication by Discord or by credential (Authentik). Once authenticated, there will be a home page with the necessary information about the project, a live page containing data of choice sent by OPC-UA and Web Socket and a final administration page in which users who have already logged in will be shown, including the possibility to change the role of each user (user or admin). This last page will only be accessible as an admin role. The last step is the XR headset part, we will also have to display some data of the automaton in the XR headset, for this we will have to use the Unity software."),
("Sensor Traineeship", "group", "More", "Sensor Traineeship", "One of the technocampus training courses that I carried out consisted of identifying sensors, analyzing their sign and searching for the operating point of the PID adjustment. I carried out tests on speed, flow and level sensors."),
("Sorting Balls", "group", "More", "Sorting Balls", "The second disciplinary project involved sorting two types of balls, yellow and pink, into three different crates: one for yellow, one for pink, and one for the others. An embedded system with sensors, motors, and other components was required to display data on a website with access to history and live data, as well as a focus on design and physical structure. After a lot of 3D modeling and research on components, my team came up with the idea of a rollercoaster. An object is thrown into the system, its color is checked by an AI camera named HuskyLens (trained to detect yellow and pink balls). If the camera detects the ball's color correctly, it directs it to the corresponding crate. If the object is unknown, or the crate is full, the ball goes to the 'other' crate. All steps are displayed on an LCD connected via I2C communication and on a website. JavaScript and Twig templates for the frontend, and Bootstrap for the design, were the main components used in the development of the website using the Express.js framework. Our website runs with Node.js. At the end of the project, we were interviewed by Télésambre, who wanted to learn more about our work. I had a lot of difficulty at the start of the project because I’m often distracted, which made teamwork challenging."),
]
for title, keyword, button, subtitle, desc in projects:
    p(title)
    p(keyword)
    p(button)
    p(subtitle)
    p("Overview")
    p(desc)
    p("Close")

h1("Internship — BlopX")
p("End-of-study placement in industrial IT: automation for a 3D print-farm and Shopify e-commerce stack. Content reflects the TFE report and marketing analysis (workflow, stack, and growth levers).")
p("Context & objectives")
p("BlopX operates as a print farm: many FDM machines, industrial\n            repeatability, and a clear product story (useful everyday objects,\n            European production, possible customization). My work centered on\n            reducing manual, repetitive steps across stock, production data,\n            catalogue, and communication — not on “doing 3D printing” for its own\n            sake, but on reliable operations and trustworthy listings.")
p("Marketing & positioning (analysis)")
p("The offer is built around coherent “universes” (outdoor,\n            equestrian, …) and creator partnerships. Growth depends on\n            discoverability and trust: SEO basics, social proof, consistent\n            creatives, and tooling that keeps publishing sustainable (scheduling,\n            repurposing, funnel thinking). Multi-channel presence (e.g. TikTok,\n            Meta, X) must stay aligned with SKU reality and production capacity.")
p("Stack & integration")
p("Core services included n8n workflows, MySQL for operational\n            data, Docker for isolated services, Home Assistant for environmental\n            supervision, and Appsmith / NocoDB to expose data to non-developers.\n            Shopify remained the commercial front; internal apps and automations\n            had to stay consistent with variants, stock movements, and fulfilment\n            events.")
p("Key deliverables")
p("Filament stock logic and movement tracking; printer telemetry\n            and command patterns (incl. MQTT / radio where applicable); pipelines\n            toward product content (structure, images, attributes) and safer SQL\n            discipline; operational dashboards; and marketing-side automation\n            drafts to keep feeds alive without manual overload.")
p("Filament stock application")
p("Mobile interface used to manage filament stock: add products, remove products, and keep the print farm inventory easier to update.")
p("Main tools used")
p("Shopify")
p("Appsmith")
p("n8n")

h1("Contact")
p("Show CV")
p("Hide CV")
p("Download")
p("Contact me")
p("Email: Bilaldrabo35@gmail.com")
p("Student account: la218086@student.helha.be")
p("LinkedIn: Bilal Drabo")
p("Phone: 0499251365")

doc.save(out)
print(out)
